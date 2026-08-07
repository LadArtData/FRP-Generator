"""Submission package assembly and export.

A package is the thing you actually submit: the narrative assembled into the
agency's required order and heading scheme, the compliance matrix, the pricing
section supplied by the approver, and the required forms, produced as a DOCX and
converted to PDF. It is versioned, downloadable, and recallable, and once
submitted or won its narrative is promoted back into the library so the next bid
starts from it.

PDF conversion runs LibreOffice headless, which is the reliable server-side path
that preserves the DOCX layout rather than re-rendering an approximation of it.
"""
from __future__ import annotations

import io
import json
import logging
import os
import shutil
import subprocess
import tempfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import audit, documents, formats, generation, opportunities, pricing
from .config import cfg
from .db import clob, cursor, transaction
from .errors import Conflict, NotFound, UpstreamError, ValidationFailed

log = logging.getLogger("harald.packages")

ACCENT = RGBColor(0x1E, 0x4B, 0x77)


# ---------------------------------------------------------------------------
# Section content assembly
# ---------------------------------------------------------------------------
def _requirement_sections(opp: dict) -> list[dict]:
    """Group drafted narrative requirements under their module headings."""
    grouped: dict[str, list[dict]] = {}
    for req in opp["requirements"]:
        if req["response_type"] != "narrative":
            continue
        body = (req["draft"] or "").strip()
        if not body:
            continue
        grouped.setdefault(req["module"], []).append(req)

    sections: list[dict] = []
    for module, title in generation.MODULE_TITLES.items():
        reqs = grouped.get(module)
        if not reqs:
            continue
        body = "\n\n".join(r["draft"].strip() for r in reqs)
        sections.append({
            "title": title,
            "body": body,
            "module_tag": module,
            "source": "generated",
            "req_ids": [r["req_id"] for r in reqs],
        })
    return sections


async def _generated_section(opp: dict, entry: dict) -> dict:
    key = entry.get("key", "")
    client = opp["client_name"] or "the client"
    briefs = {
        "transmittal": f"A letter of transmittal introducing iteria's response to "
                       f"{opp.get('solicitation_no') or 'this solicitation'} for {client}.",
        "exec_summary": f"An executive summary of iteria's Oracle Cloud Fusion proposal for {client}.",
        "qualifications": f"iteria's firm background, public-sector ERP experience, and "
                          f"relevant past performance for {client}.",
        "methodology": f"iteria's implementation methodology and phasing for {client}.",
        "project_mgmt": f"Project governance, schedule control, and risk management for {client}.",
        "staffing": f"The proposed project team and key personnel for {client}.",
        "support": f"Post go-live support, training, and service levels for {client}.",
        "references": f"References and comparable public-sector engagements relevant to {client}.",
        "contract": f"Contract alignment, service levels, and any exceptions for {client}.",
    }
    brief = briefs.get(key, f"{entry['title']} for {client}.")
    module = {"support": "TECH", "methodology": "TECH"}.get(key)
    body = await generation.draft_section(client, entry["title"], module, brief)
    return {"title": entry["title"], "body": body, "module_tag": module,
            "source": "generated", "req_ids": []}


async def assemble(opp_id: int, actor: str | None = None) -> dict:
    """Build a new package version from the current state of the bid."""
    opp = opportunities.get(opp_id)
    compliance = opportunities.compliance(opp_id)

    profile = (formats.get(opp["format_profile_id"])
               if opp["format_profile_id"] else formats.default_profile())

    sections: list[dict] = []
    for entry in profile["page_order"]:
        source = entry.get("source", "generated")

        if source == "requirements":
            requirement_sections = _requirement_sections(opp)
            if not requirement_sections:
                raise ValidationFailed(
                    "No narrative requirements have been drafted, so the proposed "
                    "solution section would be empty. Draft the requirements first."
                )
            sections.append({"title": entry["title"], "body": "", "module_tag": None,
                             "source": "generated", "req_ids": [], "heading_only": True})
            sections.extend(requirement_sections)

        elif source == "pricing":
            current = pricing.current(opp_id)
            if current:
                body = (
                    f"iteria's cost proposal for {opp['client_name']} is provided in "
                    f"{current['filename']}, submitted under separate cover as the "
                    f"solicitation requires.\n\nPricing owner: {current['owner']}. "
                    f"Version {current['version']}, status {current['status']}."
                )
            else:
                body = ("[PRICING NOT YET PROVIDED. The cost proposal is supplied by the "
                        "approver and is never generated by HARALD. This package cannot be "
                        "approved until it is attached.]")
            sections.append({"title": entry["title"], "body": body, "module_tag": None,
                             "source": "pricing", "req_ids": []})

        elif source == "form":
            forms = profile["required_forms"]
            body = ("The following forms are submitted with this response:\n"
                    + "\n".join(f"- {f}" for f in forms)) if forms else \
                   "[No required forms recorded for this agency profile.]"
            sections.append({"title": entry["title"], "body": body, "module_tag": None,
                             "source": "form", "req_ids": []})

        elif source == "manual":
            sections.append({"title": entry["title"], "body": "", "module_tag": None,
                             "source": "manual", "req_ids": []})

        else:
            sections.append(await _generated_section(opp, entry))

    with transaction() as conn:
        cur = conn.cursor()
        cur.execute("SELECT NVL(MAX(version), 0) + 1 FROM harald_packages WHERE opp_id = :o",
                    {"o": opp_id})
        version = cur.fetchone()[0]
        out = cur.var(int)
        filename = _filename(opp, version)
        cur.execute(
            """INSERT INTO harald_packages
                 (opp_id, version, status, format_profile_id, filename, compliance_json,
                  created_by)
               VALUES (:opp, :ver, 'draft', :profile, :fn, :compliance, :actor)
               RETURNING package_id INTO :out""",
            {"opp": opp_id, "ver": version, "profile": profile["profile_id"],
             "fn": filename, "compliance": json.dumps(compliance), "actor": actor,
             "out": out},
        )
        package_id = out.getvalue()[0]
        cur.executemany(
            """INSERT INTO harald_package_sections
                 (package_id, title, sort_order, body, module_tag, source, req_ids)
               VALUES (:pkg, :title, :ord, :body, :mod, :src, :reqs)""",
            [
                {"pkg": package_id, "title": s["title"], "ord": i, "body": s["body"],
                 "mod": s.get("module_tag"), "src": s["source"],
                 "reqs": json.dumps(s.get("req_ids", []))}
                for i, s in enumerate(sections)
            ],
        )

    render(package_id)
    audit.record(actor, "package.assemble", "package", package_id,
                 {"opp_id": opp_id, "version": version, "sections": len(sections)})
    log.info("assembled package_id=%s opp=%s version=%s sections=%s",
             package_id, opp_id, version, len(sections))
    return get(package_id)


def _filename(opp: dict, version: int) -> str:
    client = "".join(
        ch for ch in (opp["client_name"] or "Proposal") if ch.isalnum() or ch in " -_"
    ).strip().replace(" ", "_")
    solicitation = (opp.get("solicitation_no") or "").strip().replace(" ", "_")
    parts = ["iteria", client]
    if solicitation:
        parts.append(solicitation)
    parts.append(f"v{version}")
    return "_".join(parts) + ".docx"


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------
def _add_field(paragraph, instruction: str) -> None:
    """Insert a Word field (used for the page-number and TOC fields)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update this field in Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, placeholder, end):
        run._r.append(element)


def _style(document: Document, profile: dict) -> None:
    normal = document.styles["Normal"]
    normal.font.name = profile["font_name"]
    normal.font.size = Pt(profile["font_size"])
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), profile["font_name"])
    for section in document.sections:
        margin = Inches(profile["margin_inches"])
        section.top_margin = section.bottom_margin = margin
        section.left_margin = section.right_margin = margin


def _footer(document: Document, opp: dict) -> None:
    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = paragraph.add_run(
        f"iteria  |  {opp['client_name']}"
        f"{'  |  ' + opp['solicitation_no'] if opp.get('solicitation_no') else ''}  |  Page "
    )
    label.font.size = Pt(8)
    label.font.color.rgb = RGBColor(0x76, 0x7B, 0x83)
    _add_field(paragraph, "PAGE")


def _cover(document: Document, opp: dict, version: int) -> None:
    for _ in range(4):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(opp["title"] or "Proposal")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    for text, size in (
        (opp["client_name"] or "", 16),
        (opp.get("agency") or "", 12),
        (f"Solicitation {opp['solicitation_no']}" if opp.get("solicitation_no") else "", 12),
        (f"Due {opp['due_date']}" if opp.get("due_date") else "", 11),
    ):
        if not text:
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line = paragraph.add_run(text)
        line.font.size = Pt(size)

    for _ in range(6):
        document.add_paragraph()

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    submitted = footer.add_run(f"Submitted by iteria  |  Version {version}")
    submitted.font.size = Pt(11)
    submitted.font.color.rgb = ACCENT
    document.add_page_break()


def _toc(document: Document) -> None:
    heading = document.add_paragraph()
    run = heading.add_run("Table of Contents")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    _add_field(document.add_paragraph(), r'TOC \o "1-2" \h \z \u')
    document.add_page_break()


def _compliance_table(document: Document, opp: dict) -> None:
    heading = document.add_paragraph()
    run = heading.add_run("Requirements Compliance Matrix")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    intro = document.add_paragraph()
    intro.add_run(
        "Every requirement in the solicitation is listed below in the agency's own wording, "
        "with the section of this response that answers it."
    ).font.size = Pt(10)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for index, label in enumerate(("Ref", "Requirement", "Answered in", "Status")):
        cell = table.rows[0].cells[index]
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for req in opp["requirements"]:
        row = table.add_row().cells
        row[0].text = req["rfp_ref"] or ""
        row[1].text = req["req_text"][:400]
        row[2].text = req["section_ref"] or generation.MODULE_TITLES.get(
            req["module"], req["module"])
        row[3].text = req["status"].replace("_", " ").title()
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_page_break()


def build_docx(package_id: int) -> tuple[bytes, str]:
    package = get(package_id)
    opp = opportunities.get(package["opp_id"])
    profile = (formats.get(package["format_profile_id"])
               if package["format_profile_id"] else formats.default_profile())
    numbered = bool(profile["heading_scheme"].get("numbered", True))

    document = Document()
    _style(document, profile)
    _footer(document, opp)

    if profile["cover_required"] == "Y":
        _cover(document, opp, package["version"])
    if profile["toc_required"] == "Y":
        _toc(document)

    counter = 0
    for section in package["sections"]:
        is_subsection = section["source"] == "generated" and section["module_tag"] and \
            any(s.get("title") and s["source"] == "generated" and not s["module_tag"]
                for s in package["sections"])

        if not section["module_tag"]:
            counter += 1
            label = f"{counter}. {section['title']}" if numbered else section["title"]
            heading = document.add_heading(label, level=1)
        else:
            label = section["title"]
            heading = document.add_heading(label, level=2)
        for run in heading.runs:
            run.font.color.rgb = ACCENT

        body = (section["body"] or "").strip()
        if not body:
            note = document.add_paragraph()
            marker = note.add_run("[This section is completed by the proposal team.]")
            marker.italic = True
            marker.font.color.rgb = RGBColor(0xAB, 0x3F, 0x2C)
        else:
            for block in body.split("\n\n"):
                text = block.strip()
                if not text:
                    continue
                if text.isupper() and len(text) < 70:
                    sub = document.add_heading(text.title(), level=3)
                    for run in sub.runs:
                        run.font.color.rgb = ACCENT
                    continue
                paragraph = document.add_paragraph(text)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if section["page_break"] == "Y" and not section["module_tag"]:
            document.add_page_break()

    _compliance_table(document, opp)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read(), package["filename"]


def build_pdf(docx_bytes: bytes, filename: str) -> bytes:
    """Convert with LibreOffice headless so the PDF matches the DOCX exactly."""
    binary = shutil.which(cfg.soffice_bin) or shutil.which("libreoffice")
    if not binary:
        raise UpstreamError(
            "LibreOffice is not available in this container, so the PDF cannot be "
            "produced. The DOCX is unaffected.",
            {"looked_for": cfg.soffice_bin},
        )

    with tempfile.TemporaryDirectory() as workdir:
        source = os.path.join(workdir, filename)
        with open(source, "wb") as handle:
            handle.write(docx_bytes)
        try:
            subprocess.run(
                [binary, "--headless", "--norestore", "--convert-to", "pdf",
                 "--outdir", workdir, source],
                check=True, capture_output=True, timeout=180,
                env={**os.environ, "HOME": workdir},
            )
        except subprocess.CalledProcessError as exc:
            raise UpstreamError(
                "LibreOffice failed to convert the package to PDF.",
                {"stderr": exc.stderr.decode(errors="replace")[:500]},
            ) from exc
        except subprocess.TimeoutExpired as exc:
            raise UpstreamError("PDF conversion timed out.") from exc

        produced = os.path.join(workdir, os.path.splitext(filename)[0] + ".pdf")
        if not os.path.exists(produced):
            raise UpstreamError("LibreOffice produced no PDF output.")
        with open(produced, "rb") as handle:
            return handle.read()


def render(package_id: int) -> dict:
    """Render the package to DOCX, and to PDF when LibreOffice is present. A PDF
    failure never loses the DOCX."""
    docx_bytes, filename = build_docx(package_id)
    pdf_bytes = None
    pdf_error = None
    try:
        pdf_bytes = build_pdf(docx_bytes, filename)
    except UpstreamError as exc:
        pdf_error = exc.message
        log.warning("pdf conversion unavailable for package %s: %s", package_id, exc.message)

    with transaction() as conn:
        conn.cursor().execute(
            "UPDATE harald_packages SET docx_blob = :docx, pdf_blob = :pdf, filename = :fn "
            "WHERE package_id = :p",
            {"docx": docx_bytes, "pdf": pdf_bytes, "fn": filename, "p": package_id},
        )
    return {"package_id": package_id, "filename": filename,
            "docx_bytes": len(docx_bytes), "pdf_bytes": len(pdf_bytes) if pdf_bytes else 0,
            "pdf_error": pdf_error}


# ---------------------------------------------------------------------------
# Retrieval, lifecycle, and the compounding loop
# ---------------------------------------------------------------------------
def get(package_id: int) -> dict:
    with cursor() as cur:
        cur.execute(
            """SELECT package_id, opp_id, version, status, format_profile_id, filename,
                      compliance_json, pricing_id, approved_by, approved_at, submitted_at,
                      created_by, created_at,
                      NVL(DBMS_LOB.GETLENGTH(docx_blob), 0),
                      NVL(DBMS_LOB.GETLENGTH(pdf_blob), 0)
               FROM harald_packages WHERE package_id = :p""",
            {"p": package_id},
        )
        row = cur.fetchone()
        if not row:
            raise NotFound(f"Package {package_id} not found.")
        try:
            compliance = json.loads(clob(row[6])) if row[6] else {}
        except (json.JSONDecodeError, TypeError):
            compliance = {}
        package = {
            "package_id": row[0], "opp_id": row[1], "version": row[2], "status": row[3],
            "format_profile_id": row[4], "filename": row[5], "compliance": compliance,
            "pricing_id": row[7], "approved_by": row[8],
            "approved_at": row[9].isoformat() if row[9] else None,
            "submitted_at": row[10].isoformat() if row[10] else None,
            "created_by": row[11],
            "created_at": row[12].isoformat() if row[12] else None,
            "docx_bytes": int(row[13]), "pdf_bytes": int(row[14]), "sections": [],
        }
        cur.execute(
            """SELECT section_id, title, sort_order, body, module_tag, source, req_ids,
                      page_break
               FROM harald_package_sections WHERE package_id = :p ORDER BY sort_order""",
            {"p": package_id},
        )
        for r in cur.fetchall():
            try:
                req_ids = json.loads(clob(r[6])) if r[6] else []
            except (json.JSONDecodeError, TypeError):
                req_ids = []
            package["sections"].append({
                "section_id": r[0], "title": r[1], "sort_order": r[2], "body": clob(r[3]),
                "module_tag": r[4], "source": r[5], "req_ids": req_ids, "page_break": r[7],
            })
    return package


def list_for_opportunity(opp_id: int) -> list[dict]:
    with cursor() as cur:
        cur.execute(
            """SELECT package_id, version, status, filename, approved_by, submitted_at,
                      created_at, NVL(DBMS_LOB.GETLENGTH(pdf_blob), 0)
               FROM harald_packages WHERE opp_id = :o ORDER BY version DESC""",
            {"o": opp_id},
        )
        return [
            {"package_id": r[0], "version": r[1], "status": r[2], "filename": r[3],
             "approved_by": r[4], "submitted_at": r[5].isoformat() if r[5] else None,
             "created_at": r[6].isoformat() if r[6] else None, "has_pdf": int(r[7]) > 0}
            for r in cur.fetchall()
        ]


def update_section(section_id: int, body: str, actor: str | None = None) -> None:
    with transaction() as conn:
        cur = conn.cursor()
        cur.execute(
            "UPDATE harald_package_sections SET body = :b, source = 'manual' "
            "WHERE section_id = :s",
            {"b": body, "s": section_id},
        )
        if cur.rowcount == 0:
            raise NotFound(f"Section {section_id} not found.")
    audit.record(actor, "package.section.edit", "section", section_id)


def download(package_id: int, kind: str = "docx") -> tuple[bytes, str, str]:
    column = "pdf_blob" if kind == "pdf" else "docx_blob"
    with cursor() as cur:
        cur.execute(
            f"SELECT {column}, filename FROM harald_packages WHERE package_id = :p",
            {"p": package_id},
        )
        row = cur.fetchone()
    if not row or row[0] is None:
        raise NotFound(
            f"This package has no {kind.upper()}. Re-assemble it, or check whether PDF "
            f"conversion is available in the container."
        )
    blob = row[0].read() if hasattr(row[0], "read") else row[0]
    filename = row[1] or f"package_{package_id}.docx"
    if kind == "pdf":
        filename = os.path.splitext(filename)[0] + ".pdf"
        media = "application/pdf"
    else:
        media = ("application/vnd.openxmlformats-officedocument"
                 ".wordprocessingml.document")
    return blob, filename, media


def set_status(package_id: int, status: str, actor: str | None = None) -> None:
    if status not in ("draft", "in_review", "approved", "submitted"):
        raise ValidationFailed("Invalid package status.")
    with transaction() as conn:
        cur = conn.cursor()
        cur.execute(
            "UPDATE harald_packages SET status = :s WHERE package_id = :p",
            {"s": status, "p": package_id},
        )
        if cur.rowcount == 0:
            raise NotFound(f"Package {package_id} not found.")
    audit.record(actor, f"package.{status}", "package", package_id)


def approve(package_id: int, actor: str) -> dict:
    """Final approval. Restricted to the approver by the API layer, and blocked here
    if any mandatory requirement is still open or pricing is missing."""
    package = get(package_id)
    compliance = opportunities.compliance(package["opp_id"])
    if compliance["mandatory_gaps"] > 0:
        raise Conflict(
            f"{compliance['mandatory_gaps']} mandatory requirement(s) are not complete. "
            f"A package cannot be approved with open mandatory gaps.",
            {"mandatory_gaps": compliance["mandatory_gaps"]},
        )
    current_pricing = pricing.current(package["opp_id"])
    if not current_pricing:
        raise Conflict(
            "No pricing has been attached. The cost proposal is supplied by the approver "
            "and must be attached before the package is approved."
        )
    if current_pricing["status"] != "approved":
        raise Conflict(
            f"Pricing version {current_pricing['version']} is {current_pricing['status']}. "
            f"Mark it approved before approving the package."
        )

    with transaction() as conn:
        conn.cursor().execute(
            """UPDATE harald_packages
               SET status = 'approved', approved_by = :actor, approved_at = SYSTIMESTAMP,
                   pricing_id = :price
               WHERE package_id = :p""",
            {"actor": actor, "price": current_pricing["price_id"], "p": package_id},
        )
    audit.record(actor, "package.approve", "package", package_id,
                 {"pricing_version": current_pricing["version"]})
    return get(package_id)


def submit(package_id: int, actor: str) -> dict:
    """Mark submitted and close the compounding loop: the approved narrative becomes
    a library document, so the next bid retrieves from this one."""
    package = get(package_id)
    if package["status"] != "approved":
        raise Conflict("Only an approved package can be marked submitted.")

    docx_bytes, filename, _ = download(package_id, "docx")
    opp = opportunities.get(package["opp_id"])

    with transaction() as conn:
        conn.cursor().execute(
            """UPDATE harald_packages SET status = 'submitted', submitted_at = SYSTIMESTAMP
               WHERE package_id = :p""",
            {"p": package_id},
        )
    opportunities.update(package["opp_id"], {"status": "submitted"}, actor)

    stored = documents.store(
        filename, docx_bytes, opp_id=package["opp_id"], doc_role="iteria_response",
        doc_class="ITERIA_NARRATIVE", client_name=opp["client_name"],
        outcome="in_progress", actor=actor,
    )
    audit.record(actor, "package.submit", "package", package_id,
                 {"library_doc_id": stored["doc_id"], "chunks": stored["chunks"]})
    log.info("submitted package_id=%s and indexed %s chunks into the library",
             package_id, stored["chunks"])
    return {"package": get(package_id), "library": stored}
