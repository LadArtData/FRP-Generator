"""Render a proposal draft as a submittable Word document.

The previous export wrote a title, an italic metadata line, a bullet list of
form fields, and then the whole draft as undifferentiated paragraphs. That is a
transcript, not a proposal. An evaluator opening it sees no title page, no table
of contents, no page numbers, and no section structure — and public-sector RFPs
score "quality, clarity and completeness of proposal" as an explicit criterion.

This module renders the draft into the shape a procurement office expects:

    Title page  →  Table of contents  →  numbered sections  →  page numbers

The draft is authored in a light markdown subset because that is what the
drafting model produces and what a human can edit without a WYSIWYG editor:

    # / ## / ###   headings (mapped to Word's built-in Heading styles, so the
                   table of contents field can find them)
    | a | b |      pipe tables, rendered as real Word tables
    - / *          bullets
    1.             numbered lists
    **bold**       inline bold
    ---            horizontal rule (rendered as a paragraph border)
    > quote        block quote

Everything else is a paragraph.
"""
from __future__ import annotations

import re
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1E, 0x4B, 0x77)
MUTED = RGBColor(0x59, 0x59, 0x59)


# ---------------------------------------------------------------------------
# Low-level Word field plumbing
# ---------------------------------------------------------------------------

def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    """Insert a Word field code (TOC, PAGE, NUMPAGES).

    Fields are computed by Word, not by us, which is why the table of contents
    can carry live page numbers we have no way of knowing at render time.
    """
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _bottom_border(paragraph, size: int = 6, color: str = "1E4B77") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def _shade(cell, hex_fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")          # never "solid" — renders black
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _repeat_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _runs(paragraph, text: str, *, bold: bool = False, size: int | None = None,
          color: RGBColor | None = None) -> None:
    """Add text to a paragraph, honouring **bold** spans."""
    position = 0
    for match in _BOLD.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            run.bold = bold
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        run = paragraph.add_run(match.group(1))
        run.bold = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Document furniture
# ---------------------------------------------------------------------------

def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(16 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def _title_page(document: Document, meta: dict) -> None:
    for _ in range(4):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta.get("title") or "Proposal")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = ACCENT

    if meta.get("subtitle"):
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(meta["subtitle"])
        run.font.size = Pt(14)
        run.font.color.rgb = MUTED

    rule = document.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bottom_border(rule, size=12)

    document.add_paragraph()
    for label, value in (
        ("Prepared for", meta.get("client")),
        ("Solicitation", meta.get("solicitation")),
        ("Response due", meta.get("due_date")),
    ):
        if not value:
            continue
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = line.add_run(f"{label}: ")
        label_run.font.color.rgb = MUTED
        label_run.font.size = Pt(11)
        value_run = line.add_run(str(value))
        value_run.bold = True
        value_run.font.size = Pt(11)

    for _ in range(6):
        document.add_paragraph()

    submitted = document.add_paragraph()
    submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = submitted.add_run("Submitted by")
    run.font.color.rgb = MUTED
    run.font.size = Pt(10)

    firm = document.add_paragraph()
    firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = firm.add_run(meta.get("firm") or "iteria.us, Inc.")
    run.bold = True
    run.font.size = Pt(14)

    for line_text in (meta.get("firm_address"), meta.get("firm_contact")):
        if not line_text:
            continue
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run(line_text)
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED

    stamp = document.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stamp.add_run(meta.get("issued") or date.today().strftime("%d %B %Y"))
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _toc_page(document: Document) -> None:
    heading = document.add_paragraph()
    run = heading.add_run("Table of Contents")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = ACCENT
    _bottom_border(heading)

    body = document.add_paragraph()
    _field(body, r'TOC \o "1-3" \h \z \u',
           "Right-click and select Update Field to build the table of contents.")

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _page_footer(document: Document, label: str) -> None:
    footer = document.sections[-1].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{label}    ·    Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    _field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" of ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    _field(paragraph, "NUMPAGES", "1")
    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


# ---------------------------------------------------------------------------
# Body rendering
# ---------------------------------------------------------------------------

def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def _split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator(line: str) -> bool:
    cells = _split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(r) for r in rows)
    table = document.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, row_cells in enumerate(rows):
        cells = table.add_row().cells
        for column in range(columns):
            text = row_cells[column] if column < len(row_cells) else ""
            paragraph = cells[column].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            _runs(paragraph, text, bold=(index == 0), size=9.5)
            if index == 0:
                _shade(cells[column], "EDF1F6")
        if index == 0:
            _repeat_header(table.rows[0])
    document.add_paragraph()


def render_body(document: Document, draft: str) -> None:
    """Render the markdown-subset draft into the document."""
    lines = (draft or "").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        # Table
        if _is_table_row(stripped):
            rows: list[list[str]] = []
            while index < len(lines) and _is_table_row(lines[index].strip()):
                if not _is_separator(lines[index]):
                    rows.append(_split_row(lines[index]))
                index += 1
            _add_table(document, rows)
            continue

        # Horizontal rule
        if re.fullmatch(r"(-{3,}|_{3,}|\*{3,})", stripped):
            rule = document.add_paragraph()
            _bottom_border(rule, size=4, color="C9D3DE")
            index += 1
            continue

        # Heading
        heading = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            document.add_heading(heading.group(2).strip(), level=level)
            index += 1
            continue

        # Block quote — consecutive '>' lines are one paragraph. Rendering them
        # line by line splits inline spans: a **bold** phrase wrapped across two
        # source lines has no closing marker on either, so both render with
        # literal asterisks.
        if stripped.startswith(">"):
            quoted = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                piece = lines[index].strip().lstrip(">").strip()
                if piece:
                    quoted.append(piece)
                index += 1
            if quoted:
                quote = document.add_paragraph()
                quote.paragraph_format.left_indent = Inches(0.35)
                _runs(quote, " ".join(quoted))
                for run in quote.runs:
                    run.italic = True
                    run.font.color.rgb = MUTED
            continue

        # Bullet
        if re.match(r"[-*+]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            _runs(paragraph, re.sub(r"^[-*+]\s+", "", stripped))
            index += 1
            continue

        # Numbered
        if re.match(r"\d+[.)]\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            _runs(paragraph, re.sub(r"^\d+[.)]\s+", "", stripped))
            index += 1
            continue

        # Paragraph — join continuation lines
        buffer = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if (not nxt or _is_table_row(nxt) or nxt.startswith("#")
                    or nxt.startswith(">")
                    or re.match(r"[-*+]\s+", nxt)
                    or re.match(r"\d+[.)]\s+", nxt)
                    or re.fullmatch(r"(-{3,}|_{3,}|\*{3,})", nxt)):
                break
            buffer.append(nxt)
            index += 1
        _runs(document.add_paragraph(), " ".join(buffer))


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def build(draft: str, meta: dict) -> Document:
    """Build the full document: title page, contents, body, page numbers."""
    document = Document()

    section = document.sections[0]
    section.page_width = Inches(8.5)          # US Letter, not A4
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    _configure_styles(document)
    _title_page(document, meta)
    _toc_page(document)

    if (draft or "").strip():
        render_body(document, draft)
    else:
        document.add_heading("Draft not yet generated", level=1)
        document.add_paragraph(
            "This proposal has no draft text. Complete the intake sections and "
            "run Generate, then export again."
        )

    _page_footer(document, meta.get("footer") or meta.get("client") or "Proposal")
    return document
